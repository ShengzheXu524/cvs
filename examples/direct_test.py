#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
直接测试优化后的提示词，绕过可能导致proxies参数问题的代码
"""

import os
import sys
import json
import docx
import logging
from pathlib import Path
from anthropic import Anthropic

# 添加src目录到路径
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from src.data_organizer import DataOrganizer
from src.csv_generator import CSVGenerator

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger("direct_test")

def read_docx(file_path):
    """读取docx文件内容"""
    logger.info(f"读取文件: {file_path}")
    
    doc = docx.Document(file_path)
    
    # 读取所有段落内容
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():  # 确保不添加空白段落
            full_text.append(para.text)
    
    # 读取所有表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():  # 确保不添加空白单元格
                    row_text.append(cell.text.strip())
            if row_text:  # 确保不添加空白行
                full_text.append(" | ".join(row_text))
    
    all_text = "\n".join(full_text)
    logger.info(f"提取了 {len(all_text)} 个字符")
    
    return all_text

def extract_structured_data(client, document_text, model="claude-3-7-sonnet-20240229"):
    """使用优化后的提示词提取结构化数据，采用流式API避免超时"""
    logger.info("发送请求到Claude API（使用流式处理）...")
    
    system_prompt = """
请分析下面的考研英语真题文档，提取所有题目内容和相关信息。
请以JSON格式返回结果，确保包含所有必要信息。

# 分析任务
请分析这份考研英语真题文档，并提取以下信息：
1. 基本信息（年份、考试类型）
2. 各部分题型的原文
3. 各题目的编号、题干、选项、答案及干扰选项
4. 完形填空和特定新题型的还原后原文

# 重要说明
- 必须返回所有52道题目的完整数据！每份考研卷都有固定的52道题，即使原文中有些题目不明显，也请确保返回完整的52道题目。
- "原文"是指该题对应题型板块的试卷上的完整原文内容，不是单个题目的句子。
- 每个题型板块的原文只需在sections部分提取一次，questions部分不应包含重复的原文。
- 完形填空题尤其重要：必须在sections.cloze中提供一份完整的原文（包含所有[1], [2]等标记），questions部分的完形填空题不要包含重复的original_text或restored_text。
- 对于阅读理解、完形填空等所有题型，原文只在sections部分出现一次，所有相同题型的题目共享同一个原文。
- 每道题的"试卷答案"字段应该只包含该题的单独答案，如"A"、"B"、"C"或"D"，不要包含完整选项内容。

# 题目编号和题型映射
- 题号1-20：完形填空 (Cloze)
- 题号21-25：阅读理解 Text 1 (Reading Text 1)
- 题号26-30：阅读理解 Text 2 (Reading Text 2)
- 题号31-35：阅读理解 Text 3 (Reading Text 3)
- 题号36-40：阅读理解 Text 4 (Reading Text 4)
- 题号41-45：新题型 (New Type)
- 题号46-50：翻译 (Translation)
- 题号51：写作A (Writing Part A)
- 题号52：写作B (Writing Part B)

# 输出格式
请以下面的JSON格式返回结果：

```json
{
  "metadata": {
    "year": "年份，如2024",
    "exam_type": "考试类型，如'英语（一）'"
  },
  "sections": {
    "cloze": {
      "original_text": "完形填空原文（卷面版本，包含[1], [2]等标记，包含整篇文章）",
      "restored_text": "完形填空原文（还原后版本，已将答案填入）",
      "answers_summary": "完形填空答案汇总，如'1.D 2.C 3.B...'"
    },
    "reading": {
      "text_1": {
        "original_text": "阅读Text 1原文（整篇文章）",
        "answers_summary": "阅读Text 1答案汇总，如'21.D 22.D 23.A...'"
      },
      "text_2": {
        "original_text": "阅读Text 2原文（整篇文章）",
        "answers_summary": "阅读Text 2答案汇总"
      },
      "text_3": {
        "original_text": "阅读Text 3原文（整篇文章）",
        "answers_summary": "阅读Text 3答案汇总"
      },
      "text_4": {
        "original_text": "阅读Text 4原文（整篇文章）",
        "answers_summary": "阅读Text 4答案汇总"
      }
    },
    "new_type": {
      "original_text": "新题型原文（完整文章）",
      "restored_text": "新题型还原后原文（如果有的话，否则与原文相同）",
      "answers_summary": "新题型答案汇总"
    },
    "translation": {
      "original_text": "翻译题原文（完整段落）",
      "answers_summary": "无客观答案，可填写'N/A'"
    },
    "writing": {
      "part_a": {
        "original_text": "写作A题目（完整题目描述）",
        "answers_summary": "无客观答案，可填写'N/A'"
      },
      "part_b": {
        "original_text": "写作B题目（完整题目描述）",
        "answers_summary": "无客观答案，可填写'N/A'"
      }
    }
  },
  "questions": [
    {
      "number": 1,
      "section_type": "完形填空",
      "stem": "题干（完形填空通常为空或'/'）",
      "options": "A. Without, B. Though, C. Despite, D. Besides",
      "correct_answer": "A. Without",
      "distractor_options": "B. Though, C. Despite, D. Besides"
    },
    // ... 其他题目，必须包含所有52个题目
    {
      "number": 52,
      "section_type": "写作B",
      "stem": "写作B题目描述",
      "options": null,
      "correct_answer": null,
      "distractor_options": null
    }
  ]
}
```

注意事项：
1. sections部分要提供每个题型的完整原文，这些原文将被所有相应题号的题目共享
2. questions部分不要包含重复的原文，只包含题目的具体信息（题号、题干、选项等）
3. 对于完形填空题，原文只在sections.cloze中出现一次，不要在每道题的数据中重复包含原文

请提取完整的题型原文，确保"原文（卷面）"字段包含整篇文章或段落，而不仅仅是单个题目对应的句子。请务必返回所有52道题的完整数据，对于文档中不明确的题目，可以标记为[缺失数据]但必须保持题号的完整性。

只返回提取的JSON数据，不要包含任何其他解释或说明文字。确保JSON格式正确无误。
"""
    
    try:
        # 使用流式API
        with client.messages.stream(
            model=model,
            system=system_prompt,
            messages=[{"role": "user", "content": document_text}],
            max_tokens=30000,
            temperature=0.1,
        ) as stream:
            response_text = ""
            # 打印进度更新
            logger.info("开始接收流式响应...")
            for text in stream.text_stream:
                # 累积响应文本
                response_text += text
                # 打印进度（可选）
                if len(response_text) % 1000 == 0:
                    logger.info(f"已接收 {len(response_text)} 个字符")
            
            logger.info(f"完成，共接收 {len(response_text)} 个字符")
        
        # 保存原始响应以便检查
        with open("./test_results/api_response_raw.txt", "w", encoding="utf-8") as f:
            f.write(response_text)
            logger.info("原始响应已保存到 api_response_raw.txt")
        
        # 提取JSON
        try:
            # 查找JSON开始标记
            json_start = response_text.find('{')
            if json_start >= 0:
                # 尝试解析
                result = json.loads(response_text[json_start:])
                logger.info("成功解析JSON数据")
                return result
            else:
                logger.error("无法找到JSON数据开始标记")
                return None
        except json.JSONDecodeError as e:
            logger.error(f"JSON解析失败: {str(e)}")
            return None
    
    except Exception as e:
        logger.exception(f"API调用出错: {str(e)}")
        return None

def main():
    """主函数"""
    # 确保输出目录存在
    os.makedirs("./test_results", exist_ok=True)
    
    # 从环境变量获取API密钥
    api_key = os.getenv("CLAUDE_API_KEY")
    if not api_key:
        logger.error("未找到API密钥，请设置CLAUDE_API_KEY环境变量")
        return 1
    
    # 文件路径
    input_file = "/Users/shengzhexu/Downloads/考研英语一/2024年考研英语(一)真题及参考答案.docx"
    output_file = "./test_results/2024英语一_直接测试.csv"
    
    if not os.path.exists(input_file):
        logger.error(f"文件不存在: {input_file}")
        return 1
    
    try:
        # 读取文档内容
        document_text = read_docx(input_file)
        
        # 创建API客户端
        client = Anthropic(api_key=api_key)
        
        # 提取结构化数据
        result = extract_structured_data(client, document_text)
        
        if result:
            # 保存原始API结果
            with open("./test_results/api_result.json", "w", encoding="utf-8") as f:
                json.dump(result, f, ensure_ascii=False, indent=2)
            
            # 处理数据
            data_organizer = DataOrganizer()
            organized_data = data_organizer._process_new_format(result)
            
            # 保存处理后的数据
            with open("./test_results/organized_data.json", "w", encoding="utf-8") as f:
                json.dump(organized_data, f, ensure_ascii=False, indent=2)
            
            # 生成CSV
            csv_generator = CSVGenerator()
            if csv_generator.generate_csv(organized_data, output_file):
                logger.info(f"成功生成CSV文件: {output_file}")
                return 0
            else:
                logger.error("生成CSV文件失败")
                return 1
        else:
            logger.error("API返回的数据无效")
            return 1
            
    except Exception as e:
        logger.exception(f"处理时出错: {str(e)}")
        return 1

if __name__ == "__main__":
    sys.exit(main()) 