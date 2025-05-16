#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
文档读取模块，负责读取docx文件并预处理文本内容。
"""

import os
import re
import docx
import logging
from pathlib import Path
from docx.opc.exceptions import PackageNotFoundError

logger = logging.getLogger("考研英语真题处理.docx_reader")

class DocxReader:
    """
    DOCX文件读取器，用于读取考研英语真题文档。
    支持不同格式的考研英语真题文档，自动识别关键部分。
    """
    
    def __init__(self):
        """
        初始化文档读取器。
        """
        # 考研英语真题常见部分标记
        self.section_markers = {
            'cloze': [r'完形填空', r'Section\s+[A-Za-z].*?[完|Cloze]', r'Part\s+[A-Za-z].*?完形'],
            'reading': [r'阅读理解', r'Section\s+[A-Za-z].*?[阅读|Reading]', r'Part\s+[A-Za-z].*?阅读'],
            'new_questions': [r'新题型', r'新题型英语', r'Section\s+[A-Za-z].*?新题型'],
            'translation': [r'翻译', r'英译汉', r'Section\s+[A-Za-z].*?翻译'],
            'writing': [r'写作', r'Section\s+[A-Za-z].*?写作'],
            'text': [r'Text\s+\d+', r'Text [A-Za-z]']
        }
        
        # 答案识别模式
        self.answer_patterns = [
            r'(?:参考)?答案[与及]?解析',
            r'答案[:：]',
            r'Answer\s+Key',
            r'Keys?(\s+to)?(\s+the)?(\s+questions)?'
        ]
    
    def read_file(self, file_path):
        """
        读取docx文件并提取文本内容。
        
        Args:
            file_path (str): docx文件路径
        
        Returns:
            str: 提取的文本内容
        """
        if not os.path.exists(file_path):
            logger.error(f"文件不存在: {file_path}")
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        logger.info(f"开始读取文件: {file_path}")
        
        try:
            # 尝试从文件名提取年份信息
            file_name = os.path.basename(file_path)
            year_match = re.search(r'(\d{4})', file_name)
            year = year_match.group(1) if year_match else None
            logger.info(f"检测到年份: {year if year else '未知'}")
            
            doc = self._open_docx(file_path)
            if doc is None:
                return None
            
            # 读取所有段落内容
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():  # 确保不添加空白段落
                    full_text.append(para.text)
            
            # 读取所有表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():  # 确保不添加空白单元格
                            row_text.append(cell.text.strip())
                    if row_text:  # 确保不添加空白行
                        full_text.append(" | ".join(row_text))
            
            all_text = "\n".join(full_text)
            logger.debug(f"提取了 {len(all_text)} 个字符")
            
            # 对提取的文本进行预处理
            processed_text = self.preprocess_text(all_text)
            
            # 根据年份应用特定处理逻辑
            if year:
                processed_text = self._apply_year_specific_processing(processed_text, year)
            
            return processed_text
            
        except Exception as e:
            logger.error(f"读取文件时出错: {str(e)}")
            raise
    
    def _open_docx(self, file_path):
        """
        尝试打开docx文件，处理可能的异常。
        
        Args:
            file_path (str): docx文件路径
            
        Returns:
            docx.Document or None: 文档对象，失败时返回None
        """
        try:
            return docx.Document(file_path)
        except PackageNotFoundError:
            logger.error(f"无法打开文件，不是有效的docx格式: {file_path}")
            return None
        except Exception as e:
            logger.error(f"打开docx文件时发生错误: {str(e)}")
            return None
    
    def _apply_year_specific_processing(self, text, year):
        """
        根据年份应用特定的文本处理逻辑。
        不同年份的真题可能有不同的格式和需要特殊处理的部分。
        
        Args:
            text (str): 已提取的文本
            year (str): 年份
            
        Returns:
            str: 处理后的文本
        """
        year = int(year)
        
        # 2020年及之后的真题
        if year >= 2020:
            # 处理特定年份的格式问题
            pass
        
        # 2015-2019年的真题
        elif 2015 <= year < 2020:
            # 处理特定年份的格式问题
            pass
        
        # 2010-2014年的真题
        elif 2010 <= year < 2015:
            # 处理特定年份的格式问题
            pass
        
        # 对所有年份的通用处理
        
        # 确保题目编号格式一致
        text = re.sub(r'(?<!\d)(\d{1,2})[\s\.、]+', r'\1. ', text)
        
        # 确保选项格式一致
        text = re.sub(r'(?<![A-Za-z])(A|B|C|D)[\s\.、]+', r'[\1] ', text)
        
        return text
    
    def read_file_with_structure(self, file_path):
        """
        读取docx文件并尝试保留更多结构信息。
        
        Args:
            file_path (str): docx文件路径
        
        Returns:
            dict: 包含文档结构信息的字典
        """
        if not os.path.exists(file_path):
            logger.error(f"文件不存在: {file_path}")
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        logger.info(f"开始读取文件（保留结构）: {file_path}")
        
        try:
            doc = self._open_docx(file_path)
            if doc is None:
                return None
            
            # 提取文档基本信息
            result = {
                "title": self._extract_title(doc),
                "sections": self._extract_sections(doc),
                "paragraphs": [p.text for p in doc.paragraphs if p.text.strip()],
                "tables": self._extract_tables(doc),
                "full_text": self.read_file(file_path),
                "metadata": self._extract_metadata(doc, file_path)
            }
            
            # 尝试识别并标记文档各部分
            sections = self._identify_sections(result["full_text"])
            if sections:
                result["identified_sections"] = sections
            
            return result
            
        except Exception as e:
            logger.error(f"读取文件结构时出错: {str(e)}")
            raise
    
    def _extract_metadata(self, doc, file_path):
        """
        提取文档元数据。
        
        Args:
            doc (docx.Document): 文档对象
            file_path (str): 文件路径
            
        Returns:
            dict: 元数据字典
        """
        metadata = {}
        
        # 从文件名提取年份
        file_name = os.path.basename(file_path)
        year_match = re.search(r'(\d{4})', file_name)
        if year_match:
            metadata["year"] = year_match.group(1)
        
        # 从文件名或内容判断考试类型
        if re.search(r'英语[（\(]?一[）\)]?|英一', file_name, re.IGNORECASE):
            metadata["exam_type"] = "英语（一）"
        elif re.search(r'英语[（\(]?二[）\)]?|英二', file_name, re.IGNORECASE):
            metadata["exam_type"] = "英语（二）"
        else:
            # 从文档内容中尝试识别
            content = "\n".join([p.text for p in doc.paragraphs[:10] if p.text.strip()])
            if re.search(r'英语[（\(]?一[）\)]?|英一', content, re.IGNORECASE):
                metadata["exam_type"] = "英语（一）"
            elif re.search(r'英语[（\(]?二[）\)]?|英二', content, re.IGNORECASE):
                metadata["exam_type"] = "英语（二）"
            else:
                metadata["exam_type"] = "未知类型"
        
        return metadata
    
    def _extract_title(self, doc):
        """
        尝试从文档中提取标题。
        
        Args:
            doc (docx.Document): 文档对象
        
        Returns:
            str: 提取的标题
        """
        # 通常标题是文档的第一个段落，并且有特殊格式
        if doc.paragraphs and doc.paragraphs[0].text.strip():
            return doc.paragraphs[0].text.strip()
        return ""
    
    def _extract_sections(self, doc):
        """
        尝试从文档中提取各个部分。
        
        Args:
            doc (docx.Document): 文档对象
        
        Returns:
            dict: 各部分内容的字典
        """
        # 根据常见的题型标记识别各部分
        sections = {}
        current_section = "默认部分"
        current_content = []
        
        # 构建正则表达式来识别各部分
        section_patterns = [
            (r"Section\s+[A-D]", ""),  # 如 Section A
            (r"Part\s+[A-D]", ""),  # 如 Part A
            (r"完形填空", "完形填空"),
            (r"阅读理解", "阅读理解"),
            (r"Text\s+[1-4]", ""),  # 如 Text 1
            (r"新题型", "新题型"),
            (r"翻译", "翻译"),
            (r"写作", "写作"),
            (r"Section\s+[A-D].*?完形填空", "完形填空"),
            (r"Section\s+[A-D].*?阅读理解", "阅读理解"),
            (r"Section\s+[A-D].*?翻译", "翻译"),
            (r"Section\s+[A-D].*?写作", "写作")
        ]
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            # 检查是否是新部分的开始
            section_found = False
            for pattern, name in section_patterns:
                if re.search(pattern, text, re.IGNORECASE):
                    # 保存之前的部分内容
                    if current_content:
                        sections[current_section] = "\n".join(current_content)
                    
                    # 开始新部分
                    current_section = name if name else text
                    current_content = [text]
                    section_found = True
                    break
            
            if not section_found:
                current_content.append(text)
        
        # 保存最后一个部分
        if current_content:
            sections[current_section] = "\n".join(current_content)
            
        return sections
    
    def _identify_sections(self, text):
        """
        识别文本中的各个部分。
        
        Args:
            text (str): 文档全文
            
        Returns:
            dict: 各部分及其位置
        """
        sections = {}
        
        # 识别各部分
        for section_name, patterns in self.section_markers.items():
            for pattern in patterns:
                matches = list(re.finditer(pattern, text, re.IGNORECASE))
                if matches:
                    sections[section_name] = {
                        "start_pos": matches[0].start(),
                        "pattern_matched": pattern,
                        "text_matched": matches[0].group()
                    }
                    break
        
        # 识别答案部分
        for pattern in self.answer_patterns:
            matches = list(re.finditer(pattern, text, re.IGNORECASE))
            if matches:
                sections["answers"] = {
                    "start_pos": matches[0].start(),
                    "pattern_matched": pattern,
                    "text_matched": matches[0].group()
                }
                break
        
        # 按照位置排序
        ordered_sections = []
        for name, info in sections.items():
            ordered_sections.append((name, info["start_pos"]))
        
        ordered_sections.sort(key=lambda x: x[1])
        
        # 计算各部分的文本范围
        result = {}
        for i, (name, start_pos) in enumerate(ordered_sections):
            end_pos = len(text)
            if i < len(ordered_sections) - 1:
                end_pos = ordered_sections[i + 1][1]
            
            section_text = text[start_pos:end_pos].strip()
            result[name] = {
                "position": i + 1,
                "start_pos": start_pos,
                "end_pos": end_pos,
                "text": section_text
            }
        
        return result
    
    def _extract_tables(self, doc):
        """
        提取文档中的所有表格。
        
        Args:
            doc (docx.Document): 文档对象
        
        Returns:
            list: 表格内容列表
        """
        tables = []
        
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):  # 确保不添加完全空白的行
                    table_data.append(row_data)
            
            if table_data:  # 确保不添加完全空白的表格
                tables.append(table_data)
                
        return tables
        
    def preprocess_text(self, text):
        """
        对提取的文本进行预处理，清理不必要的内容。
        
        Args:
            text (str): 原始文本
        
        Returns:
            str: 预处理后的文本
        """
        # 删除多余的空白字符
        text = re.sub(r'\s+', ' ', text)
        
        # 清理常见的文档干扰内容
        noise_patterns = [
            r'(\d+)\s*/\s*(\d+)',  # 页码 如 "1/10"
            r'Page\s*\d+\s*of\s*\d+',  # 英文页码 如 "Page 1 of 10"
            r'考研英语网 http://.+?com',  # 网站水印
            r'[\u4e00-\u9fa5]{2,}网 https?://.+',  # 中文网站水印
            r'仅供参考.{0,10}禁止复制',  # 版权声明
            r'版权所有.{0,10}违者必究',  # 版权声明
            r'CopyRight.{0,30}Reserved',  # 英文版权声明
        ]
        
        for pattern in noise_patterns:
            text = re.sub(pattern, '', text)
        
        # 替换常见的特殊字符
        text = text.replace('–', '-')
        text = text.replace(''', "'")
        text = text.replace(''', "'")
        text = text.replace('"', '"')
        text = text.replace('"', '"')
        
        # 清理页眉页脚等内容
        text = re.sub(r'页码\s*\d+\s*', '', text)
        text = re.sub(r'Page\s*\d+\s*', '', text)
        
        # 统一空行
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # 统一完形填空空格标记
        text = re.sub(r'(__+|_\s+_|\[\s*\d+\s*\])', lambda m: f"[{len(m.group())}]", text)
        
        # 统一选项标记
        text = re.sub(r'\s([A-D])\s*[.。、]', r' [\1] ', text)
        
        return text.strip()
        
    def extract_answer_key(self, text):
        """
        从文本中提取答案部分。
        
        Args:
            text (str): 文档文本
            
        Returns:
            str: 提取的答案部分，如果未找到则返回None
        """
        # 尝试不同的答案标识
        for pattern in self.answer_patterns:
            match = re.search(f"{pattern}.*?$", text, re.DOTALL | re.IGNORECASE)
            if match:
                return match.group().strip()
        
        return None 